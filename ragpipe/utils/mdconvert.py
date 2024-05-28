
# Convert Markdown to PDF using mdpdf


import os,sys
import subprocess
# import markdown
from docxtpl import DocxTemplate, RichText 
from docx import Document
from markdown import markdown
from bs4 import BeautifulSoup
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE

def markdown_to_pdf(markdown_file_path):
    """
    Convert Markdown to PDF

    Parameters
    ----------
    markdown_file_path : str
        The path to the Markdown file
    """
    # Check if the file exists
    if not os.path.exists(markdown_file_path):
        raise FileNotFoundError(f"The file {markdown_file_path} does not exist")

    # Generate the PDF file path
    base_name = os.path.splitext(markdown_file_path)[0]
    pdf_file_path = base_name + '.pdf'

    # Convert MD to PDF and save it
    cmd = f"mdpdf -o {pdf_file_path} {markdown_file_path}"
    subprocess.run(cmd, shell=True, check=True)


def markdown_to_html(markdown_file_path):
    """
    Convert Markdown to HTML

    Parameters
    ----------
    markdown_file_path : str
        The path to the Markdown file
    """
    # Check if the file exists
    if not os.path.exists(markdown_file_path):
        raise FileNotFoundError(f"The file {markdown_file_path} does not exist")

    # Generate the PDF file path
    base_name = os.path.splitext(markdown_file_path)[0]
    html_file_path = base_name + '.html'

    # Convert MD to PDF and save it
    markdown.markdownFromFile(
        input=markdown_file_path,
        output=html_file_path,
        encoding='utf8',
        )
    

def report_to_docx(markdown_file_path,ORG_NAME,TITLE_NAME,RESEARCH_PERIOD, AUTHOR, IMPACT_PERIOD, RESEARCH_TOPIC, q1, q2, q3, q4, refs):
    """
    Convert Markdown to HTML

    Parameters
    ----------
    markdown_file_path : str
        The path to the Markdown file
    """
    # Check if the file exists
    if not os.path.exists(markdown_file_path):
        raise FileNotFoundError(f"The file {markdown_file_path} does not exist")
    
    # convert questions to docx
    converted_answers = convert_answers_to_docx([q1, q2, q3, q4])

    # Generate the PDF file path
    base_name = os.path.splitext(markdown_file_path)[0]
    docx_file_path = base_name + '.docx'

    # Convert MD to PDF and save it
    doc = DocxTemplate("./templates/RIAF_template.docx")
    docx_content = { 
                "ORG_NAME":ORG_NAME, 
                "TITLE_NAME": TITLE_NAME,
                "RESEARCH_PERIOD":RESEARCH_PERIOD,
                "AUTHOR":AUTHOR,
                "IMPACT_PERIOD":IMPACT_PERIOD,
                "RESEARCH_TOPIC":RESEARCH_TOPIC,
                "q1":converted_answers[0],
                "q2":converted_answers[1],
                "q3":converted_answers[2],
                "q4":converted_answers[4],
                "refs":refs
                }
    doc.render(docx_content)
    doc.save(docx_file_path)


def markdown_to_docx(markdown_text):
    html = markdown(markdown_text)
    soup = BeautifulSoup(html, features="html.parser")
    doc = Document()

    for element in soup:
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'h4':
            doc.add_heading(element.text, level=4)
        elif element.name == 'h5':
            doc.add_heading(element.text, level=5)
        elif element.name == 'h6':
            doc.add_heading(element.text, level=6)
        elif element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'strong':
            doc.add_paragraph(element.text, style='Bold')
        elif element.name == 'em':
            doc.add_paragraph(element.text, style='Italic')
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='ListBullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='ListNumber')

    return doc

def docx_to_richtext(doc):
    rt = RichText()
    for para in doc.paragraphs:
        for run in para.runs:
            rt.add(run.text, bold=run.bold, italic=run.italic, underline=run.underline)
        rt.add('\n')  # Add a new line after each paragraph
    return rt

def markdown_to_richtext(markdown_text):
    doc = markdown_to_docx(markdown_text)
    return docx_to_richtext(doc)


def clone_formatting(source, target):
    """ Clone formatting from source paragraph to target paragraph, including font styles and paragraph spacing. """
    target.style = source.style
    target.alignment = source.alignment
    target.paragraph_format.left_indent = source.paragraph_format.left_indent
    target.paragraph_format.right_indent = source.paragraph_format.right_indent
    target.paragraph_format.first_line_indent = source.paragraph_format.first_line_indent
    target.paragraph_format.space_before = source.paragraph_format.space_before
    target.paragraph_format.space_after = source.paragraph_format.space_after
    target.paragraph_format.line_spacing = source.paragraph_format.line_spacing
    target.paragraph_format.keep_together = source.paragraph_format.keep_together
    target.paragraph_format.keep_with_next = source.paragraph_format.keep_with_next
    target.paragraph_format.page_break_before = source.paragraph_format.page_break_before
    target.paragraph_format.widow_control = source.paragraph_format.widow_control

def clone_paragraph(p, cell):
    """Clone a paragraph into a specified table cell with full formatting, including lists."""
    target_p = cell.add_paragraph(style = p.style)
    clone_formatting(p, target_p)
    for run in p.runs:
        target_run = target_p.add_run(run.text)
        # Copy run formatting
        target_run.bold = run.bold
        target_run.italic = run.italic
        target_run.underline = run.underline
        target_run.font.size = run.font.size
        target_run.font.color.rgb = run.font.color.rgb
        target_run.font.name = run.font.name
        target_run.style = run.style

def ensure_styles(source_doc, target_doc):
    for style in source_doc.styles:
        if style.name not in [s.name for s in target_doc.styles]:
            # If the style doesn't exist, create a new one based on a built-in style if possible
            new_style = target_doc.styles.add_style(style.name, style.type)
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                new_style.base_style = target_doc.styles['Normal']  # Base new styles on 'Normal'
            # Copy style attributes as needed
            new_style.font.name = style.font.name
            new_style.font.size = style.font.size
            # Add other attributes as needed

def insert_content_into_cell(source_docx, target_docx, table_index, row_index, col_index):
    """Insert content from a source DOCX into a specific cell of a target DOCX table."""
    source_doc = Document(source_docx)
    target_doc = Document(target_docx)
    source_styles = [s for s in source_doc.styles]
    target_styles = [s for s in target_doc.styles]
    target_style_names = [s.name for s in target_styles]
    # add source styles to target doc
    for style in source_styles:
        if style.name not in target_style_names:
            target_doc.styles.add_style(style.name, style.type)
    ensure_styles(source_doc, target_doc)

    # Identify the target cell
    cell = target_doc.tables[table_index].rows[row_index].cells[col_index]

    # Clear existing cell content
    cell.text = ""
    cell.paragraphs[0].clear()

    for para in source_doc.paragraphs:
        clone_paragraph(para, cell)

    target_doc.save('updated_target.docx')


"""
def markdown_to_richtext(markdown_text):
    html = markdown(markdown_text)
    soup = BeautifulSoup(html, features="html.parser")
    rt = RichText()

    for element in soup:
        if element.name == 'h1':
            rt.add(element.text, style='Heading1')
        elif element.name == 'h2':
            rt.add(element.text, style='Heading2')
        elif element.name == 'h3':
            rt.add(element.text, style='Heading3')
        elif element.name == 'h4':
            rt.add(element.text, style='Heading4')
        elif element.name == 'h5':
            rt.add(element.text, style='Heading5')
        elif element.name == 'h6':
            rt.add(element.text, style='Heading6')
        elif element.name == 'p':
            rt.add(element.text, style='Normal')
        elif element.name == 'strong':
            rt.add(element.text, bold=True)
        elif element.name == 'em':
            rt.add(element.text, italic=True)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                rt.add(f"\n• {li.text}", style='ListBullet')
        elif element.name == 'ol':
            for idx, li in enumerate(element.find_all('li'), 1):
                rt.add(f"\n{idx}. {li.text}", style='ListNumber')

    return rt
"""

# Convert list answers to DOCX format
def convert_answers_to_docx(answers):
    docs = []
    for answer in answers:
        doc = markdown_to_docx(answer)
        docx_content = "\n".join([p.text for p in doc.paragraphs])
        docs.append(docx_content)
    return docs
    

if __name__ == '__main__':  
    if 'markdown_to_docx' in sys.argv:  
        markdown_to_docx(sys.argv[2], 
                         "an org", 
                         "a good title", 
                         "2010", 
                         "butters", 
                         "2020",
                         "how to make bread",
                         "this is the answer", 
                         "it is a good one",
                         "We should know about it",
                         "is this your final answer",
                         "This is a reference")  


def markdown_to_question(markdown_file_path, start_text, end_text):
    """
    Grabs specific questions from the output Markdown

    Parameters
    ----------
    markdown_file_path : str
        The path to the Markdown file

    # Example usage:
    start_text = "## What is the problem this research seeks to address and why is it significant? (Max 250 words)"
    end_text = "**References**"
    question_text = markdown_to_question(start_text, end_text)
    print(question_text)
    """
    # Check if the file exists
    if not os.path.exists(markdown_file_path):
        raise FileNotFoundError(f"The file {markdown_file_path} does not exist")

    question_text = ""
    capture = False

    try:
        with open(markdown_file_path, "r", encoding="utf-8") as file:
            for line in file:
                # Check if the line contains the start text and begin capture
                if start_text in line:
                    capture = True
                    continue  # Skip the line with the start text

                # Check if the line contains the end text and stop capture
                if end_text in line and capture:
                    break

                # Capture the text if between start_text and end_text
                if capture:
                    question_text += line

    except FileNotFoundError:
        print("The file 'output.md' was not found.")
    except Exception as e:
        print(f"An error occurred: {e}")

    return question_text.strip()  # Remove any leading/trailing whitespace
