# Export functions Llama-index Document files

from llama_index.core import Document
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os
import datetime
import re


def clean_text(text):
    """
    Clean text by handling various types of line breaks and spacing issues
    
    Args:
        text (str): Input text to clean
    Returns:
        str: Cleaned text
    """
    if not text:
        return ""
    
    # Replace multiple line breaks with two line breaks
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Remove extra spaces before/after line breaks
    text = re.sub(r' *\n *', '\n', text)
    
    # Remove extra spaces
    text = re.sub(r' +', ' ', text)
    
    # Handle common HTML-style breaks
    text = re.sub(r'<br\s*/?>|<p>|</p>', '\n', text)
    
    # Handle special characters often found in converted texts
    text = text.replace('\r', '\n')
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text)
    
    # Fix spacing issues
    # Fix multiple spaces
    text = re.sub(r' +', ' ', text)
    
    # Step 4: Handle paragraph breaks
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Step 5: Clean special characters
    #text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', text)
    
    # Step 6: Final formatting
    # Remove spaces at start/end of lines
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    text = text.strip()
    # Remove empty lines at start/end of document
    text = text.strip()
    
    return text




def export_documents_to_json(documents, output_file):
    """
    Export documents to a JSON file with readable format
    """
    docs_export = []
    
    for doc in documents:
        doc_dict = {
            'text': doc.text,
            'metadata': doc.metadata,
            'doc_id': doc.doc_id,
            'embedding': doc.embedding.tolist() if doc.embedding is not None else None,
            'excluded_embed_metadata_keys': doc.excluded_embed_metadata_keys,
            'excluded_llm_metadata_keys': doc.excluded_llm_metadata_keys,
            'relationships': doc.relationships,
        }
        docs_export.append(doc_dict)
    
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(docs_export, f, indent=2, ensure_ascii=False)


def export_documents_to_txt(documents, output_file):
    """
    Export documents to a text file with clear sections
    """
    with open(output_file, 'w', encoding='utf-8') as f:
        for i, doc in enumerate(documents, 1):
            f.write(f"=== Document {i} ===\n")
            f.write(f"Doc ID: {doc.doc_id}\n")
            f.write("\n=== Text ===\n")
            f.write(doc.text)
            f.write("\n\n=== Metadata ===\n")
            for key, value in doc.metadata.items():
                f.write(f"{key}: {value}\n")
            f.write("\n" + "="*50 + "\n\n")



def export_documents_to_docx_formatted(documents, outpath):
    """
    Export documents to a nicely formatted Word document
    """
    doc = DocxDocument()
    
    # Document title
    title = doc.add_heading('Document Export', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add summary
    doc.add_paragraph(f'Total Documents: {len(documents)}')
    doc.add_paragraph(f'Export Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_page_break()
    
    for i, document in enumerate(documents, 1):
        # Document header
        heading = doc.add_heading(f'Document {i}', level=1)
        heading.style.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
        
        # Document ID
        id_para = doc.add_paragraph()
        id_run = id_para.add_run(f'Document ID: {document.doc_id}')
        id_run.bold = True
        
        # Metadata section
        doc.add_heading('Metadata', level=2)
        metadata_table = doc.add_table(rows=1, cols=2)
        metadata_table.style = 'Table Grid'
        metadata_table.allow_autofit = True
        
        # Style the header row
        header_cells = metadata_table.rows[0].cells
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(11)
        header_cells[0].text = 'Key'
        header_cells[1].text = 'Value'
        
        # Add metadata rows
        for key, value in document.metadata.items():
            row_cells = metadata_table.add_row().cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)
        
        # Content section
        doc.add_heading('Content', level=2)
        content_para = doc.add_paragraph()
        content_para.add_run(document.text).font.size = Pt(11)
        
        # Add page break between documents
        doc.add_page_break()
    
    # Save the document
    output_file = os.path.join(outpath, 'exported_documents.docx')
    doc.save(output_file)



def export_document_to_docx(document, outpath):
    """
    Export a single Document object to a Word document
    
    Args:
        document: LlamaIndex Document object
        output_file (str): Path for the output .docx file
    """
    doc = DocxDocument()
    
    # Add title using document id or first few words of content
    #title = document.doc_id or document.text[:50] + "..."
    
    if document.metadata:
        try:
            title = document.metadata['title']
        except:
            title = document.text[:50] + "..."
    else:
        title = document.text[:50] + "..."
    doc.add_heading(title, 0)
    
    # Add metadata section if exists
    if document.metadata:
        doc.add_heading('Metadata', level=1)
        metadata_table = doc.add_table(rows=1, cols=2)
        metadata_table.style = 'Table Grid'
        
        # Add headers
        header_cells = metadata_table.rows[0].cells
        header_cells[0].text = 'Key'
        header_cells[1].text = 'Value'
        
        # Add metadata rows
        for key, value in document.metadata.items():
            row_cells = metadata_table.add_row().cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)
    
    # Add content section
    doc.add_heading('Content', level=1)

    # Clean and split the text into paragraphs
    cleaned_text = clean_text(document.text)
    paragraphs = cleaned_text.split('\n\n')
    
    # Add each paragraph separately for better formatting
    for para_text in paragraphs:
        if para_text.strip():  # Only add non-empty paragraphs
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(para_text.strip())
            run.font.size = Pt(11)
    #content_para = doc.add_paragraph()
    #content_para.add_run(document.text).font.size = Pt(11)
    
    # Save the document
    output_file = os.path.join(outpath, f'{title}.docx')
    doc.save(output_file)

