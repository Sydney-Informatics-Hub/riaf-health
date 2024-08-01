#Utility functions for converting Markdown to DOCX format and inserting content into DOCX files.

from docxtpl import DocxTemplate
import os
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches
import pypandoc
from docxcompose.composer import Composer
import logging

def append_doc_riaf(src_paths, 
                    dst_path, 
                    out_path = None, 
                    fname_question_titles = "./templates/question_titles.txt"):
    """
    Append the contents of one Word document to another.

    Args:
        src_paths (str): Path to the source document.
        dst_path (str): Path to the destination document.
        out_path (str): Path to the output document. If None, the destination document will be overwritten.
        fname_question_titles (str): Path to the file containing the question titles.
    """
    # check if src_paths is a list
    if not isinstance(src_paths, list):
        src_paths = [src_paths]

    if out_path is None:
        out_path = dst_path

    # Load the source and destination documents
    dst_doc = Document(dst_path)

    for i, src_path in enumerate(src_paths):
        # Load the source and destination documents
        src_doc = Document(src_path)

        def set_paragraph_styles(paragraph, text_color, bg_color, fontsize=11):
            # Set the text color and background
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(*text_color)
                run.font.size = Pt(fontsize)
            
            # Set the paragraph background color
            p = paragraph._element
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), bg_color)
            p.get_or_add_pPr().append(shading)

        # set text color and the background color of the first paragraph (title)
        #first_paragraph = src_doc.paragraphs[0] #.clear()
        try:
            questions = []
            with open(fname_question_titles, 'r') as file:
                questions = file.readlines()
            # format
            questions = [q.strip() for q in questions]
            # Add i. to the question titles and line space
            questions = [f"{i+1}. {q}" for i, q in enumerate(questions)]
            # Add empty line to dst_doc at end
            if i > 0:
                dst_doc.add_paragraph("")
            # Add the question 
            dst_doc.add_paragraph(questions[i])
            end_paragraph = dst_doc.paragraphs[-1]
            set_paragraph_styles(end_paragraph, text_color=(255, 255, 255), bg_color='000000') 
        except:
            logging.error(f"Error setting paragraph styles for question {i+1}")

        composer = Composer(dst_doc)
        composer.append(src_doc)
    composer.save(out_path)


def markdown_to_docx(input_path):
    """
    Convert a Markdown file to a DOCX file.

    Args:
        input_path (str): Path to the input Markdown file.
    """
    output_path_docx = os.path.splitext(input_path)[0] + ".docx"
    output_path_html = os.path.splitext(input_path)[0] + ".html"

    pypandoc.convert_file(input_path, format='md', to='html', outputfile=output_path_html)
    pypandoc.convert_file(output_path_html, format='html', to='docx', outputfile=output_path_docx)

    # remove the temporary HTML file
    os.remove(output_path_html)


def md2html(input_path):
    """
    Convert a Markdown file to an HTML file.

    Args:
        input_path (str): Path to the input Markdown file.
    """
    output_path = os.path.splitext(input_path)[0] + ".html"
    #pypandoc.convert_file(input_path, format='md', to='html5', outputfile=output_path, extra_args=[f"--css={css_file}"])
    pypandoc.convert_file(input_path, format='md', to='html5', outputfile=output_path)

  

def md2pdf(input_path):
    """
    Convert a Markdown file to a PDF file.

    Args:
        input_path (str): Path to the input Markdown file.
    """
    output_path = os.path.splitext(input_path)[0] + ".pdf"
    pypandoc.convert_file(input_path, format='md', to='pdf', outputfile=output_path, extra_args=['-V', 'geometry:margin=2cm'])


# Function to convert markdown to DOCX format
def markdown_to_docx2(markdown_text):
    html = markdown(markdown_text)
    soup = BeautifulSoup(html, features="html.parser")
    doc = Document()

    for element in soup:
        if element.name == 'h1':
            docdco.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'h4':
            doc.add_heading(element.text, level=4)
        elif element.name == 'h5':
            doc.add_heading(element.text, level=5)
        elif element.name == 'h6':
            doc.add_heading(element.text, level=6)
        elif element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'strong':
            doc.add_paragraph(element.text, style='Bold')
        elif element.name == 'em':
            doc.add_paragraph(element.text, style='Italic')
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='ListBullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='ListNumber')

    return doc


def clone_formatting(source, target):
    """ Clone formatting from source paragraph to target paragraph, including font styles and paragraph spacing. """
    target.style = source.style
    target.alignment = source.alignment
    target.paragraph_format.left_indent = source.paragraph_format.left_indent
    target.paragraph_format.right_indent = source.paragraph_format.right_indent
    target.paragraph_format.first_line_indent = source.paragraph_format.first_line_indent
    target.paragraph_format.space_before = source.paragraph_format.space_before
    target.paragraph_format.space_after = source.paragraph_format.space_after
    target.paragraph_format.line_spacing = source.paragraph_format.line_spacing
    target.paragraph_format.keep_together = source.paragraph_format.keep_together
    target.paragraph_format.keep_with_next = source.paragraph_format.keep_with_next
    target.paragraph_format.page_break_before = source.paragraph_format.page_break_before
    target.paragraph_format.widow_control = source.paragraph_format.widow_control

def clone_paragraph(p, cell):
    """Clone a paragraph into a specified table cell with full formatting, including lists."""
    target_p = cell.add_paragraph(style = p.style)
    clone_formatting(p, target_p)
    for run in p.runs:
        target_run = target_p.add_run(run.text)
        # Copy run formatting
        target_run.bold = run.bold
        target_run.italic = run.italic
        target_run.underline = run.underline
        target_run.font.size = run.font.size
        target_run.font.color.rgb = run.font.color.rgb
        target_run.font.name = run.font.name
        target_run.style = run.style

def ensure_styles(source_doc, target_doc):
    for style in source_doc.styles:
        if style.name not in [s.name for s in target_doc.styles]:
            # If the style doesn't exist, create a new one based on a built-in style if possible
            new_style = target_doc.styles.add_style(style.name, style.type)
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                new_style.base_style = target_doc.styles['Normal']  # Base new styles on 'Normal'
            # Copy style attributes as needed
            new_style.font.name = style.font.name
            new_style.font.size = style.font.size
            # Add other attributes as needed

def insert_content_into_cell(source_docx, target_docx, table_index, row_index, col_index):
    """Insert content from a source DOCX into a specific cell of a target DOCX table."""
    source_doc = Document(source_docx)
    target_doc = Document(target_docx)
    ensure_styles(source_doc, target_doc)

    # Identify the target cell
    cell = target_doc.tables[table_index].rows[row_index].cells[col_index]

    # Clear existing cell content
    cell.text = ""
    cell.paragraphs[0].clear()

    for para in source_doc.paragraphs:
        clone_paragraph(para, cell)

    target_doc.save('updated_target.docx')

""" Usage
source_docx = '/mnt/data/markdown_text.docx'  # Path to the source DOCX file
target_docx = 'target_template.docx'  # Path to the target DOCX template
insert_content_into_cell(source_docx, target_docx, table_index=0, row_index=1, col_index=1)
"""

def insert_content_into_placeholder(target_doc_path, source_doc_path, placeholder):
    target_doc = docx.Document(target_doc_path)
    source_doc = docx.Document(source_doc_path)

    # Search for the placeholder in target document tables
    for table in target_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    # Clear existing cell content
                    cell.text = ''
                    # Insert content from source document
                    for para in source_doc.paragraphs:
                        copy_paragraph_to_cell(para, cell)

# Convert list answers to DOCX format
def convert_answers_to_docx(answers):
    docs = []
    for answer in answers:
        doc = markdown_to_docx(answer)
        #docs.append(doc)
        docx_content = "\n".join([p.text for p in doc.paragraphs])
        docs.append(docx_content)
    return docs

def replace_html_and_convert_to_docx(html_template_path, replacements, output_docx_path, title):
    """
    Reads an HTML file, replaces placeholders with HTML content, and saves as a docx file with a specified title.
    
    Args:
        html_template_path (str): Path to the HTML template file.
        replacements (dict): Dictionary of placeholders and their corresponding HTML replacements.
        output_docx_path (str): Path where the output docx file should be saved.
        title (str): Title for the generated DOCX file.
    """
    # Read the HTML template file
    with open(html_template_path, 'r', encoding='utf-8') as file:
        html_content = file.read()
    
    # Replace placeholders in the HTML content
    for placeholder, replacement_html in replacements.items():
        # Using regular expressions to replace placeholders that include HTML content safely
        pattern = re.escape(placeholder)  # Escape to handle special characters in the placeholder
        html_content = re.sub(pattern, replacement_html, html_content)
    
    # Convert the modified HTML to a docx file
    docx_bytes_io = html2docx(html_content, title)
    
    # Save the docx file
    with open(output_docx_path, 'wb') as f:
        f.write(docx_bytes_io.getvalue())  # getvalue() extracts bytes from BytesIO object

def replace_in_html_and_convert_to_docx(docx_template_path, replacements, output_docx_path):
    """
    Convert DOCX to HTML, replace placeholders in HTML, and convert back to DOCX.
    
    Args:
        docx_template_path (str): Path to the DOCX template file.
        replacements (dict): Dictionary of placeholders and their corresponding HTML replacements.
        output_docx_path (str): Path where the output DOCX file should be saved.
    """
    # Convert DOCX to HTML
    with open(docx_template_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html_content = result.value  # The generated HTML

    # Replace placeholders in the HTML content
    for placeholder, replacement_html in replacements.items():
        pattern = re.escape(placeholder)  # Ensure the placeholder is treated as a literal string in regex
        html_content = re.sub(pattern, replacement_html, html_content)

    # Save the modified HTML to an HTML file
    with open(html_output_path, 'w', encoding='utf-8') as html_file:
        html_file.write(html_content)
        print(f"HTML file saved to {html_output_path}")

    # Convert the modified HTML back to DOCX
    pypandoc.convert_text(html_content, 'docx', format='html', outputfile=docx_output_path)
    print(f"DOCX file created at {docx_output_path}")

"""
# Example usage
replacements = {
    '{{ q1 }}': '<p>This is an answer to question 1.</p>',
    '{{ q2 }}': '<p>This is an answer to question 2.</p>',
    '{{ q3 }}': '<ul><li>Item 1 for Q3</li><li>Item 2 for Q3</li></ul>',
    '{{ q4 }}': '<p>This is an answer to question 4 with <strong>bold</strong> text.</p>'
}
replace_in_html_and_convert_to_docx('path_to_docx_template.docx', replacements, 'output_document.docx')

# Assuming the following attributes exist within your class or context
self.organisation = "Organization Name"
self.research_topic = "Research Topic"
self.research_period = "Research Period"
self.author = "Author Name"
self.impact_period = "Impact Period"
self.list_answers = [
    "### Answer to question 1 in markdown format",
    "**Answer to question 2 in markdown format**",
    "Answer to question 3 in markdown format",
    "*Answer to question 4 in markdown format*"
]
sources_text = "References text"

# Convert list answers to DOCX content
converted_answers = convert_answers_to_docx(self.list_answers)

# Create DOCX template and render content
doc = DocxTemplate("./templates/RIAF_template.docx")
docx_content = { 
    "ORG_NAME": self.organisation, 
    "TITLE_NAME": self.research_topic,
    "RESEARCH_PERIOD": self.research_period,
    "AUTHOR": self.author,
    "IMPACT_PERIOD": self.impact_period,
    "RESEARCH_TOPIC": self.research_topic,
    "q1": converted_answers[0],
    "q2": converted_answers[1],
    "q3": converted_answers[2],
    "q4": converted_answers[3],
    "refs": sources_text
}
doc.render(docx_content)
doc.save(os.path.join(self.outpath, "Use_Case_Study.docx"))

docx_content = { 
    "ORG_NAME": organisation, 
    "TITLE_NAME": research_topic,
    "RESEARCH_PERIOD": research_period,
    "AUTHOR": author,
    "IMPACT_PERIOD": impact_period,
    "RESEARCH_TOPIC": research_topic,
    "q1": docs[0],
    "q2": docs[1],
    "q3": docs[2],
    "q4": docs[3],
    "refs": sources_text
}
"""